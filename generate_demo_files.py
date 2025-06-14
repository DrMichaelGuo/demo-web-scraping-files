import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

# Simple Excel file creation without openpyxl
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime

def create_simple_xlsx(filename, data):
    """Create a simple XLSX file without openpyxl dependency"""
    # Create the basic XLSX structure
    xl_workbook = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
<sheets>
<sheet name="Sheet1" sheetId="1" r:id="rId1"/>
</sheets>
</workbook>'''

    xl_worksheet = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
<sheetData>'''

    # Add data rows
    for row_idx, row_data in enumerate(data, 1):
        xl_worksheet += f'<row r="{row_idx}">'
        for col_idx, cell_value in enumerate(row_data, 1):
            col_letter = chr(64 + col_idx)  # A, B, C
            xl_worksheet += f'<c r="{col_letter}{row_idx}" t="inlineStr"><is><t>{cell_value}</t></is></c>'
        xl_worksheet += '</row>'

    xl_worksheet += '''</sheetData>
</worksheet>'''

    xl_rels = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>
</Relationships>'''

    content_types = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
<Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
</Types>'''

    app_rels = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
</Relationships>'''

    # Create the XLSX file
    with zipfile.ZipFile(filename, 'w', zipfile.ZIP_DEFLATED) as zipf:
        zipf.writestr('[Content_Types].xml', content_types)
        zipf.writestr('_rels/.rels', app_rels)
        zipf.writestr('xl/workbook.xml', xl_workbook)
        zipf.writestr('xl/_rels/workbook.xml.rels', xl_rels)
        zipf.writestr('xl/worksheets/sheet1.xml', xl_worksheet)

def create_word_files():
    """Generate 10 Word files with title and dummy text"""
    for i in range(1, 11):
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Demo Word File {i}', 0)
        
        # Add simple one-line paragraph
        doc.add_paragraph(f"This is a demo Word document number {i}.")
        
        filename = f'downloads/demo_word_{i}.docx'
        doc.save(filename)
        print(f"Created {filename}")

def create_excel_files():
    """Generate 10 Excel files with 3x5 table of dummy data"""
    for i in range(1, 11):
        # Create 3 columns x 5 rows of dummy data
        data = []
        for row in range(1, 6):  # 5 rows
            row_data = []
            for col in range(1, 4):  # 3 columns
                cell_value = f"{chr(64 + col)}{row}"  # A1, B1, C1, A2, B2, etc.
                row_data.append(cell_value)
            data.append(row_data)
        
        filename = f'downloads/demo_excel_{i}.xlsx'
        create_simple_xlsx(filename, data)
        print(f"Created {filename}")

def create_pdf_files():
    """Generate 10 PDF files with heading and dummy text"""
    for i in range(1, 11):
        filename = f'downloads/demo_pdf_{i}.pdf'
        
        # Create PDF with SimpleDocTemplate for better text handling
        doc = SimpleDocTemplate(filename, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add heading
        heading = Paragraph(f"Demo PDF File {i}", styles['Title'])
        story.append(heading)
        story.append(Spacer(1, 12))
        
        # Add simple one-line paragraph
        text = f"This is demo PDF document number {i}."
        paragraph = Paragraph(text, styles['Normal'])
        story.append(paragraph)
        
        doc.build(story)
        print(f"Created {filename}")

def create_index_html():
    """Generate index.html with links to all files"""
    html_content = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Demo Files Index</title>
</head>
<body>
    <h1>Demo Files Index</h1>
    
    <h2>Word Documents</h2>
    <ul>"""
    
    # Add Word document links
    for i in range(1, 11):
        html_content += f'\n        <li><a href="downloads/demo_word_{i}.docx">demo_word_{i}.docx</a></li>'
    
    html_content += """
    </ul>
    
    <h2>Excel Spreadsheets</h2>
    <ul>"""
    
    # Add Excel file links
    for i in range(1, 11):
        html_content += f'\n        <li><a href="downloads/demo_excel_{i}.xlsx">demo_excel_{i}.xlsx</a></li>'
    
    html_content += """
    </ul>
    
    <h2>PDF Files</h2>
    <ul>"""
    
    # Add PDF file links
    for i in range(1, 11):
        html_content += f'\n        <li><a href="downloads/demo_pdf_{i}.pdf">demo_pdf_{i}.pdf</a></li>'
    
    html_content += """
    </ul>
</body>
</html>"""
    
    with open('index.html', 'w', encoding='utf-8') as f:
        f.write(html_content)
    print("Created index.html")

if __name__ == "__main__":
    print("Generating demo files...")
    
    # Create all file types
    create_word_files()
    create_excel_files()
    create_pdf_files()
    create_index_html()
    
    print("\nAll files generated successfully!")
    print("Open index.html in a web browser to access all the demo files.")